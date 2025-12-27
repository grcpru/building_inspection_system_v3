"""
Professional Word Report Generator for API/Database Inspections
================================================================
Based on working word_generator.py template
Adapted for database queries and SafetyCulture API photos
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from datetime import datetime
import pandas as pd
import os
import tempfile
from io import BytesIO
import re
import requests

# Safe matplotlib import
try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    plt = None

try:
    import numpy as np
    NUMPY_AVAILABLE = True
except ImportError:
    NUMPY_AVAILABLE = False
    np = None

# ═══════════════════════════════════════════════════════════════════
# METRICS CALCULATION
# ═══════════════════════════════════════════════════════════════════

def calculate_metrics(processed_data, total_inspections, building_name, address, 
                     inspection_date, inspection_date_range, is_multi_day):
    """Calculate metrics from processed data"""
    
    print("📊 Calculating metrics...")
    
    metrics = {}
    
    # Basic info
    metrics['building_name'] = building_name
    metrics['address'] = address
    metrics['inspection_date'] = inspection_date
    metrics['inspection_date_range'] = inspection_date_range
    metrics['is_multi_day_inspection'] = is_multi_day
    
    # Defect counts
    metrics['total_defects'] = len(processed_data)
    metrics['total_inspections'] = total_inspections
    metrics['total_units'] = processed_data['Unit'].nunique()
    
    print(f"   Total defects: {metrics['total_defects']}")
    print(f"   Total inspections: {metrics['total_inspections']}")
    print(f"   Total units: {metrics['total_units']}")
    
    # Defect rate
    if total_inspections > 0:
        metrics['defect_rate'] = (metrics['total_defects'] / total_inspections) * 100
    else:
        metrics['defect_rate'] = 0
    
    # Average defects per unit
    if metrics['total_units'] > 0:
        metrics['avg_defects_per_unit'] = metrics['total_defects'] / metrics['total_units']
    else:
        metrics['avg_defects_per_unit'] = 0
    
    # Unit summary by defect count
    unit_defects = processed_data.groupby('Unit').size().reset_index(name='DefectCount')
    unit_defects = unit_defects.sort_values('DefectCount', ascending=False)
    metrics['summary_unit'] = unit_defects
    
    print(f"   Unit summary shape: {unit_defects.shape}")
    
    # Unit categories
    metrics['ready_units'] = len(unit_defects[unit_defects['DefectCount'] <= 2])
    metrics['minor_work_units'] = len(unit_defects[(unit_defects['DefectCount'] >= 3) & (unit_defects['DefectCount'] <= 7)])
    metrics['major_work_units'] = len(unit_defects[(unit_defects['DefectCount'] >= 8) & (unit_defects['DefectCount'] <= 14)])
    metrics['extensive_work_units'] = len(unit_defects[unit_defects['DefectCount'] >= 15])
    
    # Percentages
    total_units = metrics['total_units']
    if total_units > 0:
        metrics['ready_pct'] = (metrics['ready_units'] / total_units) * 100
        metrics['minor_pct'] = (metrics['minor_work_units'] / total_units) * 100
        metrics['major_pct'] = (metrics['major_work_units'] / total_units) * 100
        metrics['extensive_pct'] = (metrics['extensive_work_units'] / total_units) * 100
    else:
        metrics['ready_pct'] = metrics['minor_pct'] = metrics['major_pct'] = metrics['extensive_pct'] = 0
    
    # Trade summary
    trade_defects = processed_data.groupby('Trade').size().reset_index(name='DefectCount')
    trade_defects = trade_defects.sort_values('DefectCount', ascending=False)
    metrics['summary_trade'] = trade_defects
    
    print(f"   Trade summary shape: {trade_defects.shape}")
    print("✅ Metrics calculated")
    
    return metrics

def create_word_report_from_database(
    inspection_ids: list,
    db_connection,
    api_key: str,
    output_path: str,
    report_type: str = "single",
    images: dict = None
) -> bool:
    """
    Generate professional Word report from database
    
    Args:
        inspection_ids: List of inspection IDs
        db_connection: Database connection (psycopg2)
        api_key: SafetyCulture API key (for photos)
        output_path: Where to save the report
        report_type: "single" or "multi" 
        images: Optional dict with 'logo' and 'cover' paths
    
    Returns:
        True if successful
    """
    
    try:
        print(f"🎨 Generating Word report (type: {report_type})...")
        
        # DEBUG: Check images
        if images:
            print(f"📷 Images provided: {images}")
            if images.get('logo'):
                print(f"   Logo path: {images['logo']} (exists: {os.path.exists(images['logo'])})")
            if images.get('cover'):
                print(f"   Cover path: {images['cover']} (exists: {os.path.exists(images['cover'])})")
        else:
            print("⚠️  No images provided")
        
        cursor = db_connection.cursor()
        
        # Get building info and dates
        cursor.execute("""
            SELECT 
                b.name as building_name,
                b.address,
                MIN(i.inspection_date) as start_date,
                MAX(i.inspection_date) as end_date,
                COUNT(DISTINCT i.id) as inspection_count
            FROM inspector_inspections i
            JOIN inspector_buildings b ON i.building_id = b.id
            WHERE i.id = ANY(%s)
            GROUP BY b.name, b.address
        """, (inspection_ids,))
        
        row = cursor.fetchone()
        if not row:
            print("❌ No building data found")
            return False
        
        building_name = row[0]
        address = row[1] or "Address not specified"
        start_date = row[2]
        end_date = row[3]
        inspection_count = row[4]
        
        # Determine inspection date display
        if start_date == end_date:
            inspection_date = start_date.strftime('%Y-%m-%d')
            inspection_date_range = inspection_date
            is_multi_day = False
        else:
            inspection_date = start_date.strftime('%Y-%m-%d')
            inspection_date_range = f"{start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')}"
            is_multi_day = True
        
        print(f"📍 Building: {building_name}")
        print(f"📅 Date: {inspection_date_range}")
        print(f"🔢 Inspections: {inspection_count}")
        
        # Get all defects with proper column names
        cursor.execute("""
            SELECT 
                ii.room,
                ii.component,
                ii.notes,
                ii.trade,
                ii.urgency,
                ii.status_class,
                ii.photo_url,
                ii.photo_media_id,
                ii.inspector_notes,
                ii.inspection_date,
                ii.unit,
                ii.unit_type
            FROM inspector_inspection_items ii
            WHERE ii.inspection_id = ANY(%s)
            AND LOWER(ii.status_class) = 'not ok'
            ORDER BY ii.unit, ii.room, ii.component
        """, (inspection_ids,))
        
        defect_rows = cursor.fetchall()
        
        # Get total inspections (all items, not just defects)
        cursor.execute("""
            SELECT COUNT(*)
            FROM inspector_inspection_items
            WHERE inspection_id = ANY(%s)
        """, (inspection_ids,))
        
        total_inspections = cursor.fetchone()[0]
        
        cursor.close()
        
        if len(defect_rows) == 0:
            print("❌ No defects found")
            return False
        
        # Convert to DataFrame with CORRECT column names
        processed_data = pd.DataFrame(defect_rows, columns=[
            'Room', 'Component', 'Issue', 'Trade', 'Severity', 'StatusClass',
            'photo_url', 'photo_media_id', 'inspector_notes', 'inspection_date',
            'Unit', 'unit_type'
        ])
        
        # CRITICAL: Set StatusClass to "Not OK" for filtering
        processed_data['StatusClass'] = 'Not OK'
        
        print(f"📊 Found {len(processed_data)} defects")
        print(f"📋 Columns: {list(processed_data.columns)}")
        print(f"🔍 Sample data:\n{processed_data.head()}")
        
        # Calculate metrics
        metrics = calculate_metrics(processed_data, total_inspections, building_name, 
                                    address, inspection_date, inspection_date_range, is_multi_day)
        
        # Generate appropriate report type
        if report_type == "single" or len(inspection_ids) == 1:
            print("📄 Generating SINGLE inspection report (detailed with photos)...")
            doc = generate_single_inspection_report(processed_data, metrics, api_key, images)
        else:
            print("📊 Generating BUILDING SUMMARY report (executive overview)...")
            doc = generate_professional_word_report(processed_data, metrics, images)
        
        # Save
        doc.save(output_path)
        print(f"✅ Saved: {output_path}")
        
        return True
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False

# ═══════════════════════════════════════════════════════════════════
# SINGLE INSPECTION REPORT - MODERN REDESIGN WITH PHOTOS
# ═══════════════════════════════════════════════════════════════════

def generate_single_inspection_report(processed_data, metrics, api_key, images=None):
    """
    Generate modern single inspection report with prominent photos
    Completely different from building summary report
    """
    
    try:
        doc = Document()
        setup_document_formatting(doc)
        
        # Add logo and cover with SINGLE UNIT styling
        if images:
            add_logo_to_header(doc, images)
        
        # Custom cover page for single unit
        add_single_unit_cover_page(doc, processed_data, metrics, images)
        
        # Unit snapshot
        add_unit_snapshot(doc, processed_data, metrics)
        
        # Room-by-room defect breakdown with photos
        add_room_by_room_defects(doc, processed_data, api_key)
        
        # Summary and recommendations
        add_single_unit_summary(doc, processed_data, metrics)
        
        print("✅ Single unit report completed!")
        return doc
    
    except Exception as e:
        print(f"❌ Error in single report: {e}")
        import traceback
        traceback.print_exc()
        return create_error_document(e, metrics)


def add_single_unit_cover_page(doc, processed_data, metrics, images=None):
    """Custom cover page for single unit inspection"""
    
    try:
        # Main title
        title_para = doc.add_paragraph()
        title_para.style = 'CleanTitle'
        title_run = title_para.add_run("UNIT INSPECTION REPORT")
        title_run.font.size = Pt(28)
        
        # Line separator
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run("─" * 40)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(12)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Get unit info
        unit = processed_data['Unit'].iloc[0] if len(processed_data) > 0 else 'Unknown'
        
        # Unit number - large and prominent
        doc.add_paragraph()
        unit_para = doc.add_paragraph()
        unit_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        unit_run = unit_para.add_run(f"UNIT {sanitize_text(str(unit))}")
        unit_run.font.name = 'Arial'
        unit_run.font.size = Pt(24)
        unit_run.font.bold = True
        unit_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Building name
        doc.add_paragraph()
        building_para = doc.add_paragraph()
        building_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        building_run = building_para.add_run(sanitize_text(metrics['building_name']))
        building_run.font.name = 'Arial'
        building_run.font.size = Pt(18)
        building_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Cover image
        if images and images.get('cover') and os.path.exists(images['cover']):
            try:
                doc.add_paragraph()
                cover_para = doc.add_paragraph()
                cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cover_para.add_run().add_picture(images['cover'], width=Inches(4.7))
                doc.add_paragraph()
            except Exception as e:
                print(f"Error adding cover image: {e}")
        
        # Quick stats box
        doc.add_paragraph()
        stats_para = doc.add_paragraph()
        stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        total_defects = len(processed_data)
        inspection_date = processed_data['inspection_date'].iloc[0] if 'inspection_date' in processed_data.columns else 'N/A'
        
        if hasattr(inspection_date, 'strftime'):
            date_str = inspection_date.strftime('%d %B %Y')
        else:
            date_str = str(inspection_date)
        
        stats_text = f"""Inspection Date: {date_str}
Total Defects Found: {total_defects}
Status: {'REQUIRES ATTENTION' if total_defects > 5 else 'MINOR WORK NEEDED'}"""
        
        stats_run = stats_para.add_run(stats_text)
        stats_run.font.name = 'Arial'
        stats_run.font.size = Pt(12)
        stats_run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in single unit cover: {e}")


def add_unit_snapshot(doc, processed_data, metrics):
    """Quick snapshot of unit condition"""
    
    try:
        header = doc.add_paragraph("UNIT SNAPSHOT")
        header.style = 'CleanSectionHeader'
        
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_para.add_run("─" * 63)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Get stats
        total_defects = len(processed_data)
        severity_counts = processed_data['Severity'].value_counts()
        trade_counts = processed_data['Trade'].value_counts()
        room_counts = processed_data['Room'].value_counts()
        
        # Create summary boxes
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        for col in table.columns:
            col.width = Inches(3.5)
        
        # Box 1: Severity Breakdown
        cell1 = table.cell(0, 0)
        set_cell_background_color(cell1, "FFE6E6")
        para1 = cell1.paragraphs[0]
        para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para1.paragraph_format.space_before = Pt(10)
        para1.paragraph_format.space_after = Pt(10)
        
        run1a = para1.add_run("SEVERITY BREAKDOWN\n")
        run1a.font.name = 'Arial'
        run1a.font.size = Pt(11)
        run1a.font.bold = True
        run1a.font.color.rgb = RGBColor(0, 0, 0)
        
        for severity, count in severity_counts.items():
            run1b = para1.add_run(f"{severity}: {count}\n")
            run1b.font.name = 'Arial'
            run1b.font.size = Pt(10)
            run1b.font.color.rgb = RGBColor(0, 0, 0)
        
        # Box 2: Top Trade
        cell2 = table.cell(0, 1)
        set_cell_background_color(cell2, "E6F3FF")
        para2 = cell2.paragraphs[0]
        para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para2.paragraph_format.space_before = Pt(10)
        para2.paragraph_format.space_after = Pt(10)
        
        run2a = para2.add_run("PRIMARY TRADE\n")
        run2a.font.name = 'Arial'
        run2a.font.size = Pt(11)
        run2a.font.bold = True
        run2a.font.color.rgb = RGBColor(0, 0, 0)
        
        if len(trade_counts) > 0:
            run2b = para2.add_run(f"{sanitize_text(trade_counts.index[0])}\n")
            run2b.font.name = 'Arial'
            run2b.font.size = Pt(14)
            run2b.font.bold = True
            run2b.font.color.rgb = RGBColor(0, 0, 0)
            
            run2c = para2.add_run(f"{trade_counts.iloc[0]} defects")
            run2c.font.name = 'Arial'
            run2c.font.size = Pt(10)
            run2c.font.color.rgb = RGBColor(0, 0, 0)
        
        # Box 3: Rooms Affected
        cell3 = table.cell(1, 0)
        set_cell_background_color(cell3, "FFF9E6")
        para3 = cell3.paragraphs[0]
        para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para3.paragraph_format.space_before = Pt(10)
        para3.paragraph_format.space_after = Pt(10)
        
        run3a = para3.add_run("ROOMS AFFECTED\n")
        run3a.font.name = 'Arial'
        run3a.font.size = Pt(11)
        run3a.font.bold = True
        run3a.font.color.rgb = RGBColor(0, 0, 0)
        
        run3b = para3.add_run(f"{len(room_counts)}\n")
        run3b.font.name = 'Arial'
        run3b.font.size = Pt(20)
        run3b.font.bold = True
        run3b.font.color.rgb = RGBColor(0, 0, 0)
        
        run3c = para3.add_run("room areas")
        run3c.font.name = 'Arial'
        run3c.font.size = Pt(10)
        run3c.font.color.rgb = RGBColor(0, 0, 0)
        
        # Box 4: Total Defects
        cell4 = table.cell(1, 1)
        set_cell_background_color(cell4, "E6FFE6")
        para4 = cell4.paragraphs[0]
        para4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para4.paragraph_format.space_before = Pt(10)
        para4.paragraph_format.space_after = Pt(10)
        
        run4a = para4.add_run("TOTAL DEFECTS\n")
        run4a.font.name = 'Arial'
        run4a.font.size = Pt(11)
        run4a.font.bold = True
        run4a.font.color.rgb = RGBColor(0, 0, 0)
        
        run4b = para4.add_run(f"{total_defects}\n")
        run4b.font.name = 'Arial'
        run4b.font.size = Pt(20)
        run4b.font.bold = True
        run4b.font.color.rgb = RGBColor(192, 0, 0) if total_defects > 10 else RGBColor(255, 140, 0)
        
        run4c = para4.add_run("items identified")
        run4c.font.name = 'Arial'
        run4c.font.size = Pt(10)
        run4c.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in unit snapshot: {e}")


def add_room_by_room_defects(doc, processed_data, api_key):
    """
    DETAILED DEFECTS - Table format matching screenshot
    Replaces the room-by-room grouping
    """
    
    try:
        # Header - "DETAILED DEFECTS" instead of "DEFECTS BY ROOM"
        header = doc.add_paragraph("DETAILED DEFECTS")
        header.style = 'CleanSectionHeader'
        
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_para.add_run("─" * 63)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()
        
        total_defects = len(processed_data)
        print(f"   Processing {total_defects} defects for detailed view...")
        
        # Process EACH defect (not grouped by room)
        for idx, (_, defect) in enumerate(processed_data.iterrows(), 1):
            print(f"   📋 Processing defect {idx} of {total_defects}...")
            
            # Defect number header
            defect_num_para = doc.add_paragraph()
            defect_num_run = defect_num_para.add_run(f"Defect {idx} of {total_defects}")
            defect_num_run.font.name = 'Arial'
            defect_num_run.font.size = Pt(11)
            defect_num_run.font.bold = True
            defect_num_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Create table - EXACTLY like screenshot
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Table Grid'
            
            table.columns[0].width = Inches(2.0)
            table.columns[1].width = Inches(4.5)
            
            # Row 1: Room/Location
            cell_label_0 = table.cell(0, 0)
            cell_value_0 = table.cell(0, 1)
            set_cell_background_color(cell_label_0, "D9D9D9")
            
            cell_label_0.text = "Room/Location"
            cell_label_0.paragraphs[0].runs[0].font.name = 'Arial'
            cell_label_0.paragraphs[0].runs[0].font.size = Pt(10)
            cell_label_0.paragraphs[0].runs[0].font.bold = True
            cell_label_0.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            
            cell_value_0.text = sanitize_text(str(defect.get('Room', 'Unknown')))
            cell_value_0.paragraphs[0].runs[0].font.name = 'Arial'
            cell_value_0.paragraphs[0].runs[0].font.size = Pt(10)
            cell_value_0.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            
            # Row 2: Component
            cell_label_1 = table.cell(1, 0)
            cell_value_1 = table.cell(1, 1)
            set_cell_background_color(cell_label_1, "D9D9D9")
            
            cell_label_1.text = "Component"
            cell_label_1.paragraphs[0].runs[0].font.name = 'Arial'
            cell_label_1.paragraphs[0].runs[0].font.size = Pt(10)
            cell_label_1.paragraphs[0].runs[0].font.bold = True
            cell_label_1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            
            cell_value_1.text = sanitize_text(str(defect.get('Component', 'Unknown')))
            cell_value_1.paragraphs[0].runs[0].font.name = 'Arial'
            cell_value_1.paragraphs[0].runs[0].font.size = Pt(10)
            cell_value_1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            
            # Row 3: Trade Category
            cell_label_2 = table.cell(2, 0)
            cell_value_2 = table.cell(2, 1)
            set_cell_background_color(cell_label_2, "D9D9D9")
            
            cell_label_2.text = "Trade Category"
            cell_label_2.paragraphs[0].runs[0].font.name = 'Arial'
            cell_label_2.paragraphs[0].runs[0].font.size = Pt(10)
            cell_label_2.paragraphs[0].runs[0].font.bold = True
            cell_label_2.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            
            cell_value_2.text = sanitize_text(str(defect.get('Trade', 'Unknown')))
            cell_value_2.paragraphs[0].runs[0].font.name = 'Arial'
            cell_value_2.paragraphs[0].runs[0].font.size = Pt(10)
            cell_value_2.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            
            # Row 4: Inspector Notes
            cell_label_3 = table.cell(3, 0)
            cell_value_3 = table.cell(3, 1)
            set_cell_background_color(cell_label_3, "D9D9D9")
            
            cell_label_3.text = "Inspector Notes"
            cell_label_3.paragraphs[0].runs[0].font.name = 'Arial'
            cell_label_3.paragraphs[0].runs[0].font.size = Pt(10)
            cell_label_3.paragraphs[0].runs[0].font.bold = True
            cell_label_3.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            
            # Get inspector notes or issue
            notes = defect.get('inspector_notes', defect.get('Issue', 'No notes'))
            if pd.isna(notes) or str(notes).strip() == '' or str(notes).lower() == 'nan':
                notes = defect.get('Issue', 'No notes')
            
            cell_value_3.text = sanitize_text(str(notes))
            cell_value_3.paragraphs[0].runs[0].font.name = 'Arial'
            cell_value_3.paragraphs[0].runs[0].font.size = Pt(10)
            cell_value_3.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            
            # Row 5: Photo Defect
            cell_label_4 = table.cell(4, 0)
            cell_value_4 = table.cell(4, 1)
            set_cell_background_color(cell_label_4, "D9D9D9")
            
            cell_label_4.text = "Photo Defect"
            cell_label_4.paragraphs[0].runs[0].font.name = 'Arial'
            cell_label_4.paragraphs[0].runs[0].font.size = Pt(10)
            cell_label_4.paragraphs[0].runs[0].font.bold = True
            cell_label_4.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
            
            # Add photo in the cell
            photo_url = defect.get('photo_url')
            print(f"      Photo URL: {photo_url}")
            
            if photo_url and str(photo_url).strip() and str(photo_url) != 'nan' and api_key:
                print(f"      📸 Downloading photo...")
                photo_data = download_photo(photo_url, api_key)
                
                if photo_data:
                    try:
                        # Clear the cell
                        cell_value_4.text = ""
                        
                        # Get the first paragraph
                        photo_para = cell_value_4.paragraphs[0]
                        photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add photo
                        run = photo_para.add_run()
                        run.add_picture(photo_data, width=Inches(4.0))
                        
                        # Add timestamp below photo
                        timestamp_para = cell_value_4.add_paragraph()
                        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        timestamp_para.paragraph_format.space_before = Pt(0)
                        timestamp_para.paragraph_format.space_after = Pt(0)
                        
                        # Format timestamp
                        inspection_date = defect.get('inspection_date', datetime.now())
                        if hasattr(inspection_date, 'strftime'):
                            timestamp_str = inspection_date.strftime('%d %b %Y at %I:%M%p').lower()
                        else:
                            try:
                                from datetime import datetime as dt
                                date_obj = dt.strptime(str(inspection_date), '%Y-%m-%d %H:%M:%S')
                                timestamp_str = date_obj.strftime('%d %b %Y at %I:%M%p').lower()
                            except:
                                timestamp_str = str(inspection_date)
                        
                        timestamp_run = timestamp_para.add_run(timestamp_str)
                        timestamp_run.font.name = 'Arial'
                        timestamp_run.font.size = Pt(10)
                        timestamp_run.font.color.rgb = RGBColor(255, 255, 255)
                        timestamp_run.font.bold = True
                        
                        print(f"      ✅ Photo added successfully")
                    except Exception as e:
                        print(f"      ❌ Error embedding photo: {e}")
                        import traceback
                        traceback.print_exc()
                        cell_value_4.text = "Photo error"
                else:
                    print(f"      ⚠️ No photo data received")
                    cell_value_4.text = "Photo not available"
            else:
                if not photo_url or str(photo_url).strip() == '' or str(photo_url) == 'nan':
                    print(f"      ⚠️ No photo URL")
                    cell_value_4.text = "No photo URL"
                elif not api_key:
                    print(f"      ⚠️ No API key")
                    cell_value_4.text = "No API key provided"
                else:
                    print(f"      ⚠️ Photo not available")
                    cell_value_4.text = "No photo"
            
            # Space between defects
            doc.add_paragraph()
            doc.add_paragraph()
        
        print(f"   ✅ Completed all {total_defects} defects")
        doc.add_page_break()
    
    except Exception as e:
        print(f"❌ Error in detailed defects: {e}")
        import traceback
        traceback.print_exc()

def add_single_unit_summary(doc, processed_data, metrics):
    """Final summary and action items for single unit"""
    
    try:
        header = doc.add_paragraph("SUMMARY & NEXT STEPS")
        header.style = 'CleanSectionHeader'
        
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_para.add_run("─" * 63)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        total_defects = len(processed_data)
        severity_counts = processed_data['Severity'].value_counts()
        trade_counts = processed_data['Trade'].value_counts()
        
        # Summary box
        summary_para = doc.add_paragraph()
        summary_para.style = 'CleanBody'
        
        summary_text = f"""**INSPECTION SUMMARY**

This unit inspection identified {total_defects} defects requiring attention before settlement. """
        
        if len(severity_counts) > 0:
            urgent = severity_counts.get('Urgent', 0)
            high = severity_counts.get('High Priority', 0)
            
            if urgent > 0:
                summary_text += f"**{urgent} urgent defect{'s' if urgent != 1 else ''}** require immediate remediation. "
            if high > 0:
                summary_text += f"{high} high priority item{'s' if high != 1 else ''} should be addressed promptly. "
        
        summary_text += f"""

**PRIMARY TRADE CATEGORIES**:
"""
        
        for trade, count in trade_counts.head(3).items():
            summary_text += f"• {sanitize_text(trade)}: {count} defect{'s' if count != 1 else ''}\n"
        
        summary_text += """
**RECOMMENDED ACTIONS**:

1. **Immediate**: Address all urgent and high priority defects within 7 days
2. **Trade Coordination**: Schedule remediation with qualified contractors
3. **Re-inspection**: Book follow-up inspection after repairs completed
4. **Documentation**: Maintain photographic evidence of all completed work
5. **Settlement**: Obtain final sign-off before proceeding to settlement

**TIMELINE**: Estimated 2-3 weeks for complete remediation based on defect complexity."""
        
        add_formatted_text_with_bold(summary_para, summary_text)
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph()
        closing_para = doc.add_paragraph()
        closing_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        closing_run = closing_para.add_run("─── END OF UNIT INSPECTION REPORT ───")
        closing_run.font.name = 'Arial'
        closing_run.font.size = Pt(12)
        closing_run.font.color.rgb = RGBColor(0, 0, 0)
        closing_run.font.bold = True
    
    except Exception as e:
        print(f"Error in unit summary: {e}")
# ═══════════════════════════════════════════════════════════════════
# REST OF THE FUNCTIONS FROM WORKING TEMPLATE
# ═══════════════════════════════════════════════════════════════════

def generate_professional_word_report(processed_data, metrics, images=None):
    """Generate professional Word report - MAIN FUNCTION"""
    
    try:
        doc = Document()
        setup_document_formatting(doc)
        add_logo_to_header(doc, images)
        add_clean_cover_page(doc, metrics, images)
        add_executive_overview(doc, metrics)
        add_inspection_process(doc, metrics)
        add_units_analysis(doc, metrics)
        add_defects_analysis(doc, processed_data, metrics)
        add_data_visualization(doc, processed_data, metrics)
        add_trade_summary(doc, processed_data, metrics)
        add_component_breakdown(doc, processed_data, metrics)
        add_recommendations(doc, metrics)
        add_footer(doc, metrics)
        
        print("✅ Report generation completed!")
        return doc
    
    except Exception as e:
        print(f"❌ Error in report generation: {e}")
        return create_error_document(e, metrics)


def setup_document_formatting(doc):
    """Setup Arial font and clean styling"""
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    styles = doc.styles
    
    # Title style
    if 'CleanTitle' not in [s.name for s in styles]:
        title_style = styles.add_style('CleanTitle', 1)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(28)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 0, 0)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        title_style.paragraph_format.space_before = Pt(10)
    
    # Section header
    if 'CleanSectionHeader' not in [s.name for s in styles]:
        section_style = styles.add_style('CleanSectionHeader', 1)
        section_font = section_style.font
        section_font.name = 'Arial'
        section_font.size = Pt(18)
        section_font.bold = True
        section_font.color.rgb = RGBColor(0, 0, 0)
        section_style.paragraph_format.space_before = Pt(20)
        section_style.paragraph_format.space_after = Pt(10)
    
    # Subsection header
    if 'CleanSubsectionHeader' not in [s.name for s in styles]:
        subsection_style = styles.add_style('CleanSubsectionHeader', 1)
        subsection_font = subsection_style.font
        subsection_font.name = 'Arial'
        subsection_font.size = Pt(14)
        subsection_font.bold = True
        subsection_font.color.rgb = RGBColor(0, 0, 0)
        subsection_style.paragraph_format.space_before = Pt(16)
        subsection_style.paragraph_format.space_after = Pt(8)
    
    # Body text
    if 'CleanBody' not in [s.name for s in styles]:
        body_style = styles.add_style('CleanBody', 1)
        body_font = body_style.font
        body_font.name = 'Arial'
        body_font.size = Pt(11)
        body_font.color.rgb = RGBColor(0, 0, 0)
        body_style.paragraph_format.line_spacing = 1.2
        body_style.paragraph_format.space_after = Pt(6)


def sanitize_text(text):
    """Sanitize text - NO HTML entities, just remove control characters"""
    if not text or text is None:
        return ""
    
    text = str(text).strip()
    
    # Remove problematic characters but KEEP & < > etc
    text = text.replace('\x00', '')
    text = text.replace('\x0b', '')
    text = text.replace('\x0c', '')
    text = text.replace('\x1f', '')
    
    text = ''.join(char for char in text if char.isprintable() or char in '\n\r\t')
    
    return text


def add_formatted_text_with_bold(paragraph, text, style_name='CleanBody'):
    """Add text with **bold** formatting support"""
    
    try:
        parts = re.split(r'\*\*(.*?)\*\*', text)
        
        for i, part in enumerate(parts):
            if i % 2 == 0:  # Regular text
                if part:
                    run = paragraph.add_run(sanitize_text(part))
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:  # Bold text
                run = paragraph.add_run(sanitize_text(part))
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.bold = True
        
        paragraph.style = style_name
    
    except Exception as e:
        run = paragraph.add_run(sanitize_text(text))
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        paragraph.style = style_name


def add_logo_to_header(doc, images=None):
    """Add logo to header"""
    
    try:
        if images and images.get('logo') and os.path.exists(images['logo']):
            section = doc.sections[0]
            header = section.header
            header.paragraphs[0].clear()
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            header_run = header_para.add_run()
            header_run.add_picture(images['logo'], width=Inches(2.0))
    except Exception as e:
        print(f"Warning: Could not add logo: {e}")


def add_clean_cover_page(doc, metrics, images=None):
    """Professional cover page"""
    
    try:
        # Title - 2 lines
        title_para = doc.add_paragraph()
        title_para.style = 'CleanTitle'
        title_run = title_para.add_run("PRE-SETTLEMENT\nINSPECTION REPORT")
        title_run.font.size = Pt(28)
        
        # Line separator
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run("─" * 40)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(12)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Building name
        doc.add_paragraph()
        building_para = doc.add_paragraph()
        building_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        building_run = building_para.add_run(sanitize_text(metrics['building_name'].upper()))
        building_run.font.name = 'Arial'
        building_run.font.size = Pt(22)
        building_run.font.bold = True
        building_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Address
        doc.add_paragraph()
        address_para = doc.add_paragraph()
        address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_run = address_para.add_run(sanitize_text(metrics['address']))
        address_run.font.name = 'Arial'
        address_run.font.size = Pt(14)
        address_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Cover image
        if images and images.get('cover') and os.path.exists(images['cover']):
            try:
                doc.add_paragraph()
                cover_para = doc.add_paragraph()
                cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cover_para.add_run().add_picture(images['cover'], width=Inches(4.7))
                doc.add_paragraph()
            except Exception as e:
                print(f"Error adding cover image: {e}")
        
        # Inspection Overview header
        doc.add_paragraph()
        overview_header = doc.add_paragraph()
        overview_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        overview_run = overview_header.add_run("INSPECTION OVERVIEW")
        overview_run.font.name = 'Arial'
        overview_run.font.size = Pt(20)
        overview_run.font.bold = True
        overview_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Line
        line_para2 = doc.add_paragraph()
        line_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run2 = line_para2.add_run("─" * 50)
        line_run2.font.name = 'Arial'
        line_run2.font.size = Pt(12)
        line_run2.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()
        
        # Metrics table
        add_metrics_table(doc, metrics)
        
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Report details - bottom left
        details_para = doc.add_paragraph()
        details_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Format dates
        if metrics.get('is_multi_day_inspection'):
            date_range = metrics['inspection_date_range']
            try:
                dates = date_range.split(' to ')
                if len(dates) == 2:
                    start_date = datetime.strptime(dates[0].strip(), '%Y-%m-%d')
                    end_date = datetime.strptime(dates[1].strip(), '%Y-%m-%d')
                    inspection_date_display = f"{start_date.strftime('%d %B %Y')} to {end_date.strftime('%d %B %Y')}"
                else:
                    inspection_date_display = date_range
            except:
                inspection_date_display = date_range
        else:
            try:
                single_date = datetime.strptime(metrics['inspection_date'], '%Y-%m-%d')
                inspection_date_display = single_date.strftime('%d %B %Y')
            except:
                inspection_date_display = metrics['inspection_date']
        
        quality_score = max(0, 100 - metrics.get('defect_rate', 0))
        
        details_text = f"""Generated on {datetime.now().strftime('%d %B %Y')}

Inspection Date: {inspection_date_display}
Units Inspected: {metrics.get('total_units', 0):,}
Components Evaluated: {metrics.get('total_inspections', 0):,}
Quality Score: {quality_score:.1f}/100"""
        
        details_run = details_para.add_run(sanitize_text(details_text))
        details_run.font.name = 'Arial'
        details_run.font.size = Pt(11)
        details_run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in cover page: {e}")


def add_metrics_table(doc, metrics):
    """Colored metrics boxes"""
    
    try:
        doc.add_paragraph()
        
        table = doc.add_table(rows=2, cols=3)
        table.alignment = 1  # Center
        
        for col in table.columns:
            col.width = Inches(2.4)
        
        metrics_data = [
            ("TOTAL UNITS", f"{metrics.get('total_units', 0):,}", "Units Inspected", "A8D3E6"),
            ("DEFECTS FOUND", f"{metrics.get('total_defects', 0):,}", f"{metrics.get('defect_rate', 0):.1f}% Rate", "FEDBDB"),
            ("READY UNITS", f"{metrics.get('ready_units', 0)}", f"{metrics.get('ready_pct', 0):.1f}%", "C8E6C9"),
            ("MINOR WORK", f"{metrics.get('minor_work_units', 0)}", f"{metrics.get('minor_pct', 0):.1f}%", "FFF3D7"),
            ("MAJOR WORK", f"{metrics.get('major_work_units', 0)}", f"{metrics.get('major_pct', 0):.1f}%", "F4C2A1"),
            ("EXTENSIVE WORK", f"{metrics.get('extensive_work_units', 0)}", f"{metrics.get('extensive_pct', 0):.1f}%", "F4A6A6")
        ]
        
        for i, (label, value, subtitle, bg_color) in enumerate(metrics_data):
            row_idx = i // 3
            col_idx = i % 3
            
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            
            # Set background
            try:
                shading_elm = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{bg_color}"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
            except:
                pass
            
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
            
            # Label
            label_run = para.add_run(f"{label}\n")
            label_run.font.name = 'Arial'
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Value
            value_run = para.add_run(f"{value}\n")
            value_run.font.name = 'Arial'
            value_run.font.size = Pt(24)
            value_run.font.bold = True
            value_run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Subtitle
            subtitle_run = para.add_run(subtitle)
            subtitle_run.font.name = 'Arial'
            subtitle_run.font.size = Pt(9)
            subtitle_run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()
    
    except Exception as e:
        print(f"Error in metrics table: {e}")


# Copy ALL remaining functions from the template...
# [I'll include the most critical ones here, but the file is getting long]

def add_executive_overview(doc, metrics):
    """Executive overview section"""
    
    try:
        header = doc.add_paragraph("EXECUTIVE OVERVIEW")
        header.style = 'CleanSectionHeader'
        
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_para.add_run("─" * 63)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Format dates
        if metrics.get('is_multi_day_inspection'):
            date_range = metrics['inspection_date_range']
            try:
                dates = date_range.split(' to ')
                if len(dates) == 2:
                    start_date = datetime.strptime(dates[0].strip(), '%Y-%m-%d')
                    end_date = datetime.strptime(dates[1].strip(), '%Y-%m-%d')
                    inspection_date_text = f"between {start_date.strftime('%d %B %Y')} and {end_date.strftime('%d %B %Y')}"
                else:
                    inspection_date_text = f"between {date_range}"
            except:
                inspection_date_text = f"between {date_range}"
        else:
            try:
                single_date = datetime.strptime(metrics['inspection_date'], '%Y-%m-%d')
                inspection_date_text = f"on {single_date.strftime('%d %B %Y')}"
            except:
                inspection_date_text = f"on {metrics['inspection_date']}"
        
        overview_text = f"""This comprehensive quality assessment encompasses the systematic evaluation of {metrics.get('total_units', 0):,} residential units within {sanitize_text(metrics.get('building_name', 'the building complex'))}, conducted {inspection_date_text}. This report was compiled on {datetime.now().strftime('%d %B %Y')}.

**Inspection Methodology**: Each unit underwent thorough room-by-room evaluation covering all major building components, including structural elements, mechanical systems, finishes, fixtures, and fittings. The assessment follows industry-standard protocols for pre-settlement quality verification.

**Key Findings**: The inspection revealed {metrics.get('total_defects', 0):,} individual defects across {metrics.get('total_inspections', 0):,} evaluated components, yielding an overall defect rate of {metrics.get('defect_rate', 0):.2f}%. Defect level analysis indicates {metrics.get('ready_pct', 0):.1f}% of units ({metrics.get('ready_units', 0)} units) require only minor work for handover.

**Strategic Insights**: The data reveals systematic patterns across trade categories, with concentrated defect types requiring targeted remediation strategies. This analysis enables optimized resource allocation and realistic timeline planning for completion preparation."""
        
        overview_para = doc.add_paragraph()
        add_formatted_text_with_bold(overview_para, overview_text)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in executive overview: {e}")


def add_inspection_process(doc, metrics):
    """Inspection process section"""
    
    try:
        header = doc.add_paragraph("INSPECTION PROCESS & METHODOLOGY")
        header.style = 'CleanSectionHeader'
        
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_para.add_run("─" * 63)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        scope_header = doc.add_paragraph("INSPECTION SCOPE & STANDARDS")
        scope_header.style = 'CleanSubsectionHeader'
        
        scope_text = f"""The comprehensive pre-settlement quality assessment was systematically executed across all {metrics.get('total_units', 0):,} residential units, encompassing detailed evaluation of {metrics.get('total_inspections', 0):,} individual components and building systems.

**Structural Assessment**
- Building envelope integrity and weatherproofing
- Structural elements and load-bearing components
- Foundation and concrete work evaluation

**Systems Evaluation**
- Electrical installations, fixtures, and safety compliance
- Plumbing systems, water pressure, and drainage
- HVAC systems and ventilation adequacy

**Finishes & Fixtures**
- Wall, ceiling, and flooring finish quality
- Door and window installation and operation
- Kitchen and bathroom fixture functionality
- Built-in storage and joinery craftsmanship"""
        
        scope_para = doc.add_paragraph()
        add_formatted_text_with_bold(scope_para, scope_text)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in inspection process: {e}")


def add_units_analysis(doc, metrics):
    """Units analysis section"""
    
    try:
        header = doc.add_paragraph("UNITS REQUIRING PRIORITY ATTENTION")
        header.style = 'CleanSectionHeader'
        
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_para.add_run("─" * 63)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        if 'summary_unit' in metrics and len(metrics['summary_unit']) > 0:
            # Create chart
            create_units_chart(doc, metrics)
            
            # Analysis text
            top_unit = metrics['summary_unit'].iloc[0]
            
            summary_text = f"""**Priority Analysis Results**: Unit {top_unit['Unit']} requires immediate priority attention with {top_unit['DefectCount']} identified defects, representing the highest concentration of remediation needs within the development.

**Resource Allocation Framework**:
- **Critical Priority**: {len(metrics['summary_unit'][metrics['summary_unit']['DefectCount'] > 15])} units requiring extensive remediation (15+ defects each)
- **High Priority**: {len(metrics['summary_unit'][(metrics['summary_unit']['DefectCount'] > 7) & (metrics['summary_unit']['DefectCount'] <= 15)])} units requiring major work (8-15 defects each)
- **Medium Priority**: {len(metrics['summary_unit'][(metrics['summary_unit']['DefectCount'] > 2) & (metrics['summary_unit']['DefectCount'] <= 7)])} units requiring intermediate work (3-7 defects each)
- **Handover Ready**: {len(metrics['summary_unit'][metrics['summary_unit']['DefectCount'] <= 2])} units ready for immediate handover

**Strategic Insights**: This distribution pattern enables targeted resource deployment and realistic timeline forecasting for completion preparation activities."""
            
            summary_para = doc.add_paragraph()
            add_formatted_text_with_bold(summary_para, summary_text)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in units analysis: {e}")


def add_defects_analysis(doc, processed_data, metrics):
    """Defects analysis section"""
    
    try:
        header = doc.add_paragraph("DEFECT PATTERNS & ANALYSIS")
        header.style = 'CleanSectionHeader'
        
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_para.add_run("─" * 63)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        if 'summary_trade' in metrics and len(metrics['summary_trade']) > 0:
            top_trade = metrics['summary_trade'].iloc[0]
            total_defects = metrics.get('total_defects', 0)
            trade_percentage = (top_trade['DefectCount']/total_defects*100) if total_defects > 0 else 0
            
            defects_text = f"""**Primary Defect Category Analysis**: The comprehensive evaluation of {total_defects:,} individually documented defects reveals "{sanitize_text(top_trade['Trade'])}" as the dominant concern category, accounting for {top_trade['DefectCount']} instances ({trade_percentage:.1f}% of total defects).

**Pattern Recognition**: This concentration within the {sanitize_text(top_trade['Trade']).lower()} trade category encompasses multiple sub-issues including installation inconsistencies, finish quality variations, functional defects, and compliance gaps.

**Strategic Implications**: The clustering of defects within specific trade categories suggests that focused remediation efforts targeting the top 3-4 trade categories could address approximately 60-80% of all identified issues."""
            
            defects_para = doc.add_paragraph()
            add_formatted_text_with_bold(defects_para, defects_text)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in defects analysis: {e}")


def add_data_visualization(doc, processed_data, metrics):
    """Data visualization section"""
    
    try:
        header = doc.add_paragraph("COMPREHENSIVE DATA VISUALISATION")
        header.style = 'CleanSectionHeader'
        
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_para.add_run("─" * 63)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        intro_text = "This section presents visual analytics of the inspection data, highlighting key patterns and trends to support strategic decision-making and resource allocation."
        intro_para = doc.add_paragraph(intro_text)
        intro_para.style = 'CleanBody'
        
        doc.add_paragraph()
        
        # Create charts
        create_pie_chart(doc, metrics)
        create_severity_chart(doc, metrics)
        create_trade_chart(doc, metrics)
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in data visualization: {e}")


def add_trade_summary(doc, processed_data, metrics):
    """Trade summary section - FIXED WITH DEBUG"""
    
    try:
        header = doc.add_paragraph("TRADE-SPECIFIC DEFECT ANALYSIS")
        header.style = 'CleanSectionHeader'
        
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_para.add_run("─" * 63)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        overview_text = """This section provides a comprehensive breakdown of identified defects organized by trade category, including complete unit inventories for targeted remediation planning and resource allocation optimization."""
        
        overview_para = doc.add_paragraph(overview_text)
        overview_para.style = 'CleanBody'
        
        doc.add_paragraph()
        
        # DEBUG
        print(f"\n🔍 TRADE SUMMARY DEBUG:")
        print(f"   processed_data shape: {processed_data.shape}")
        print(f"   processed_data columns: {list(processed_data.columns)}")
        print(f"   StatusClass values: {processed_data['StatusClass'].unique()}")
        print(f"   Sample data:\n{processed_data[['Trade', 'Component', 'Unit', 'StatusClass']].head()}")
        
        if processed_data is not None and len(processed_data) > 0:
            try:
                component_details = generate_complete_component_details(processed_data)
                
                print(f"   Component details shape: {component_details.shape}")
                
                if len(component_details) > 0:
                    print(f"   ✅ Adding trade tables with {len(component_details)} rows")
                    add_trade_tables(doc, component_details)
                else:
                    print(f"   ⚠️ No component details generated")
                    # Add placeholder
                    no_data = doc.add_paragraph("No trade-specific data available.")
                    no_data.style = 'CleanBody'
            except Exception as e:
                print(f"   ❌ Error generating trade tables: {e}")
                import traceback
                traceback.print_exc()
        else:
            print(f"   ⚠️ No processed_data available")
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"❌ Error in trade summary: {e}")
        import traceback
        traceback.print_exc()


def add_component_breakdown(doc, processed_data, metrics):
    """Component breakdown section - FIXED WITH DEBUG"""
    
    try:
        header = doc.add_paragraph("COMPONENT-LEVEL ANALYSIS")
        header.style = 'CleanSectionHeader'
        
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_para.add_run("─" * 63)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        intro_text = "This analysis identifies the most frequently affected individual components across all units, enabling targeted quality control improvements and preventive measures for future construction phases."
        
        intro_para = doc.add_paragraph(intro_text)
        intro_para.style = 'CleanBody'
        
        doc.add_paragraph()
        
        # DEBUG
        print(f"\n🔍 COMPONENT BREAKDOWN DEBUG:")
        print(f"   processed_data is None: {processed_data is None}")
        
        if processed_data is not None and len(processed_data) > 0:
            print(f"   processed_data length: {len(processed_data)}")
            
            component_data = generate_component_breakdown(processed_data)
            
            print(f"   component_data length: {len(component_data)}")
            
            if len(component_data) > 0:
                most_freq_header = doc.add_paragraph("Most Frequently Affected Components")
                most_freq_header.style = 'CleanSubsectionHeader'
                
                top_components = component_data.head(15)
                
                print(f"   Creating table with {len(top_components)} rows")
                
                comp_table = doc.add_table(rows=1, cols=4)
                comp_table.style = 'Table Grid'
                
                comp_table.columns[0].width = Inches(2.5)
                comp_table.columns[1].width = Inches(2.0)
                comp_table.columns[2].width = Inches(2.5)
                comp_table.columns[3].width = Inches(1.0)
                
                # Headers
                headers = ['Component', 'Trade', 'Affected Units', 'Count']
                for i, header in enumerate(headers):
                    cell = comp_table.cell(0, i)
                    cell.text = header
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.runs[0]
                    run.font.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    set_cell_background_color(cell, "F0F0F0")
                
                # Data rows
                for idx, (_, comp_row) in enumerate(top_components.iterrows()):
                    row = comp_table.add_row()
                    row_color = "FFFFFF" if idx % 2 == 0 else "F8F8F8"
                    
                    # Component
                    cell1 = row.cells[0]
                    cell1.text = sanitize_text(str(comp_row.get('Component', 'N/A')))
                    cell1.paragraphs[0].runs[0].font.name = 'Arial'
                    cell1.paragraphs[0].runs[0].font.size = Pt(10)
                    cell1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    set_cell_background_color(cell1, row_color)
                    
                    # Trade
                    cell2 = row.cells[1]
                    cell2.text = sanitize_text(str(comp_row.get('Trade', 'N/A')))
                    cell2.paragraphs[0].runs[0].font.name = 'Arial'
                    cell2.paragraphs[0].runs[0].font.size = Pt(10)
                    cell2.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    set_cell_background_color(cell2, row_color)
                    
                    # Affected units
                    units_text = sanitize_text(str(comp_row.get('Affected_Units', '')))
                    if len(units_text) > 50:
                        units_text = units_text[:47] + "..."
                    
                    cell3 = row.cells[2]
                    cell3.text = units_text
                    cell3.paragraphs[0].runs[0].font.name = 'Arial'
                    cell3.paragraphs[0].runs[0].font.size = Pt(9)
                    cell3.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    set_cell_background_color(cell3, row_color)
                    
                    # Count
                    cell4 = row.cells[3]
                    cell4.text = str(comp_row.get('Unit_Count', 0))
                    cell4.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell4.paragraphs[0].runs[0].font.name = 'Arial'
                    cell4.paragraphs[0].runs[0].font.size = Pt(10)
                    cell4.paragraphs[0].runs[0].font.bold = True
                    cell4.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    set_cell_background_color(cell4, row_color)
                
                print(f"   ✅ Table created successfully")
            else:
                print(f"   ⚠️ No component data generated")
                no_data = doc.add_paragraph("No component data available.")
                no_data.style = 'CleanBody'
        else:
            print(f"   ⚠️ No processed_data available")
        
        doc.add_page_break()
    
    except Exception as e:
        print(f"❌ Error in component breakdown: {e}")
        import traceback
        traceback.print_exc()

def add_recommendations(doc, metrics):
    """Recommendations section"""
    
    try:
        header = doc.add_paragraph("STRATEGIC RECOMMENDATIONS & ACTION PLAN")
        header.style = 'CleanSectionHeader'
        
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_para.add_run("─" * 63)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        priorities_header = doc.add_paragraph("IMMEDIATE PRIORITIES")
        priorities_header.style = 'CleanSubsectionHeader'
        
        priorities = []
        ready_pct = metrics.get('ready_pct', 0)
        extensive_units = metrics.get('extensive_work_units', 0)
        
        if ready_pct > 75:
            priorities.append("**Accelerated Completion Protocol**: With 75%+ units requiring only minor work, implement immediate handover for compliant units while establishing parallel remediation workflows for remaining inventory.")
        elif ready_pct > 50:
            priorities.append("**Phased Completion Strategy**: Establish structured completion phases prioritizing ready units first, with clear milestone-based progression for units under remediation.")
        else:
            priorities.append("**Quality-First Approach**: Implement comprehensive remediation program before handover to ensure optimal customer satisfaction and minimize post-handover defect claims.")
        
        if 'summary_trade' in metrics and len(metrics['summary_trade']) > 0:
            top_trade = metrics['summary_trade'].iloc[0]
            top_trade_pct = (top_trade['DefectCount'] / metrics.get('total_defects', 1) * 100)
            priorities.append(f"**{sanitize_text(top_trade['Trade'])} Focus Initiative**: This trade represents {top_trade_pct:.1f}% of all defects ({top_trade['DefectCount']} instances). Deploy dedicated supervision teams and additional resources with daily progress monitoring.")
        
        if extensive_units > 0:
            priorities.append(f"**Specialized Remediation Teams**: {extensive_units} units require extensive work (15+ defects each). Establish dedicated teams with enhanced supervision to maintain project timeline integrity and quality standards.")
        
        priorities.append("**Enhanced Quality Protocols**: Implement multi-tier inspection checkpoints with supervisor sign-offs for critical trades before final handover, reducing post-handover callback rates.")
        
        for i, priority in enumerate(priorities, 1):
            priority_para = doc.add_paragraph()
            add_formatted_text_with_bold(priority_para, f"{i}. {priority}")
            priority_para.paragraph_format.left_indent = Inches(0.4)
        
        doc.add_paragraph()
        doc.add_page_break()
    
    except Exception as e:
        print(f"Error in recommendations: {e}")


def add_footer(doc, metrics):
    """Footer section"""
    
    try:
        header = doc.add_paragraph("REPORT DOCUMENTATION & APPENDICES")
        header.style = 'CleanSectionHeader'
        
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_para.add_run("─" * 63)
        line_run.font.name = 'Arial'
        line_run.font.size = Pt(10)
        line_run.font.color.rgb = RGBColor(0, 0, 0)
        
        data_summary_header = doc.add_paragraph("COMPREHENSIVE INSPECTION METRICS")
        data_summary_header.style = 'CleanSubsectionHeader'
        
        avg_defects = metrics.get('avg_defects_per_unit', 0)
        defect_rate = metrics.get('defect_rate', 0)
        quality_score = max(0, 100 - defect_rate)
        
        data_summary_text = f"""**INSPECTION SCOPE & RESULTS**:
- Total Residential Units Evaluated: {metrics.get('total_units', 0):,}
- Total Building Components Assessed: {metrics.get('total_inspections', 0):,}
- Total Defects Documented: {metrics.get('total_defects', 0):,}
- Overall Defect Rate: {metrics.get('defect_rate', 0):.2f}%
- Average Defects per Unit: {avg_defects:.2f}
- Development Quality Score: {quality_score:.1f}/100

**DEFECT LEVEL FRAMEWORK DISTRIBUTION**:
- Minor Work Required: {metrics.get('ready_units', 0)} units ({metrics.get('ready_pct', 0):.1f}%)
- Intermediate Remediation Required: {metrics.get('minor_work_units', 0)} units ({metrics.get('minor_pct', 0):.1f}%)
- Major Remediation Required: {metrics.get('major_work_units', 0)} units ({metrics.get('major_pct', 0):.1f}%)
- Extensive Remediation Required: {metrics.get('extensive_work_units', 0)} units ({metrics.get('extensive_pct', 0):.1f}%)"""
        
        data_summary_para = doc.add_paragraph()
        add_formatted_text_with_bold(data_summary_para, data_summary_text)
        
        doc.add_paragraph()
        
        details_header = doc.add_paragraph("REPORT GENERATION & COMPANION RESOURCES")
        details_header.style = 'CleanSubsectionHeader'
        
        # Format dates
        if metrics.get('is_multi_day_inspection'):
            date_range = metrics['inspection_date_range']
            try:
                dates = date_range.split(' to ')
                if len(dates) == 2:
                    start_date = datetime.strptime(dates[0].strip(), '%Y-%m-%d')
                    end_date = datetime.strptime(dates[1].strip(), '%Y-%m-%d')
                    inspection_date_display = f"{start_date.strftime('%d %B %Y')} to {end_date.strftime('%d %B %Y')}"
                else:
                    inspection_date_display = date_range
            except:
                inspection_date_display = date_range
        else:
            try:
                single_date = datetime.strptime(metrics['inspection_date'], '%Y-%m-%d')
                inspection_date_display = single_date.strftime('%d %B %Y')
            except:
                inspection_date_display = metrics['inspection_date']
        
        details_text = f"""**REPORT METADATA**:
- Report Generated: {datetime.now().strftime('%d %B %Y at %I:%M %p')}
- Inspection Completion: {inspection_date_display}
- Building Development: {sanitize_text(metrics.get('building_name', 'N/A'))}
- Property Location: {sanitize_text(metrics.get('address', 'N/A'))}

**COMPANION DOCUMENTATION SUITE**:
Complete defect inventories, unit-by-unit detailed breakdowns, interactive filterable data tables, and comprehensive photographic documentation are available in the accompanying Excel analytics workbook.

**TECHNICAL SUPPORT & FOLLOW-UP**:
For technical inquiries, data interpretation assistance, or additional analysis requirements, please contact the inspection team. Ongoing support is available for remediation planning, progress tracking, and post-completion verification inspections."""
        
        details_para = doc.add_paragraph()
        add_formatted_text_with_bold(details_para, details_text)
        
        doc.add_paragraph()
        closing_para = doc.add_paragraph()
        closing_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        closing_run = closing_para.add_run("END OF REPORT")
        closing_run.font.name = 'Arial'
        closing_run.font.size = Pt(14)
        closing_run.font.color.rgb = RGBColor(0, 0, 0)
        closing_run.font.bold = True
    
    except Exception as e:
        print(f"Error in footer: {e}")

# ═══════════════════════════════════════════════════════════════════
# CHART GENERATION FUNCTIONS
# ═══════════════════════════════════════════════════════════════════

def create_units_chart(doc, metrics):
    """Create units chart with color coding"""
    
    if not MATPLOTLIB_AVAILABLE:
        add_text_units_summary(doc, metrics)
        return
    
    try:
        chart_title = doc.add_paragraph("Top 20 Units Requiring Immediate Intervention")
        chart_title.style = 'CleanSubsectionHeader'
        
        top_units = metrics['summary_unit'].head(20)
        
        if len(top_units) > 0:
            fig, ax = plt.subplots(figsize=(16, 12))
            
            # Color coding
            colors = []
            for count in top_units['DefectCount']:
                if count > 25:
                    colors.append('#ff9999')  # Critical
                elif count >= 15:
                    colors.append('#ffcc99')  # Extensive
                elif count >= 8:
                    colors.append('#ffff99')  # Major
                elif count >= 3:
                    colors.append('#99ff99')  # Minor
                else:
                    colors.append('#99ccff')  # Ready
            
            if NUMPY_AVAILABLE:
                y_pos = np.arange(len(top_units))
            else:
                y_pos = list(range(len(top_units)))
            
            bars = ax.barh(y_pos, top_units['DefectCount'], color=colors, alpha=0.8)
            
            ax.set_yticks(y_pos)
            ax.set_yticklabels([f"Unit {sanitize_text(str(unit))}" for unit in top_units['Unit']], fontsize=14)
            ax.set_xlabel('Number of Defects', fontsize=16, fontweight='600')
            ax.set_title('Units Ranked by Defect Concentration (Priority Order)',
                        fontsize=18, fontweight='600', pad=25)
            
            ax.grid(axis='x', alpha=0.3, linestyle=':')
            
            # Value labels
            for i, (bar, value) in enumerate(zip(bars, top_units['DefectCount'])):
                ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                       f'{value}', va='center', fontweight='bold', fontsize=12)
            
            # Legend
            from matplotlib.patches import Patch
            legend_elements = [
                Patch(facecolor='#ff9999', label='Critical (25+ defects)', alpha=0.8),
                Patch(facecolor='#ffcc99', label='Extensive (15-24 defects)', alpha=0.8),
                Patch(facecolor='#ffff99', label='Major (8-14 defects)', alpha=0.8),
                Patch(facecolor='#99ff99', label='Minor (3-7 defects)', alpha=0.8),
                Patch(facecolor='#99ccff', label='Ready (0-2 defects)', alpha=0.8)
            ]
            ax.legend(handles=legend_elements, loc='upper right', fontsize=14, framealpha=0.9)
            
            plt.tight_layout()
            add_chart_to_document(doc, fig)
            plt.close()
    
    except Exception as e:
        print(f"Error creating units chart: {e}")


def create_pie_chart(doc, metrics):
    """Create trade distribution pie chart"""
    
    if not MATPLOTLIB_AVAILABLE:
        add_text_trade_summary(doc, metrics)
        return
    
    try:
        if 'summary_trade' not in metrics or len(metrics['summary_trade']) == 0:
            return
        
        breakdown_header = doc.add_paragraph("Defects Distribution by Trade Category")
        breakdown_header.style = 'CleanSubsectionHeader'
        
        trade_data = metrics['summary_trade'].copy()
        total_defects = metrics.get('total_defects', 0)
        
        num_trades = len(trade_data)
        
        if NUMPY_AVAILABLE and num_trades <= 12:
            colors = plt.cm.Set3(np.linspace(0, 1, num_trades))
        else:
            base_colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99', '#ff99cc', 
                          '#c2c2f0', '#ffb3e6', '#c4e17f', '#76d7c4', '#f7dc6f']
            colors = (base_colors * ((num_trades // len(base_colors)) + 1))[:num_trades]
        
        fig, ax = plt.subplots(figsize=(10, 8))
        
        # Sanitize trade names
        trade_labels = [sanitize_text(str(trade)) for trade in trade_data['Trade']]
        
        wedges, texts, autotexts = ax.pie(
            trade_data['DefectCount'], 
            labels=trade_labels, 
            colors=colors,
            autopct='%1.1f%%',
            startangle=45
        )
        
        ax.set_title(f'Distribution of Defects by Trade Category ({num_trades} Trades)', 
                    fontsize=16, fontweight='600', pad=20)
        
        plt.tight_layout()
        add_chart_to_document(doc, fig)
        plt.close()
        
        # Summary text
        if len(trade_data) > 0:
            top_trade = trade_data.iloc[0]
            summary_text = f"""The analysis reveals {sanitize_text(top_trade['Trade'])} as the primary defect category, representing {top_trade['DefectCount']} of the total {total_defects:,} defects ({top_trade['DefectCount']/total_defects*100:.1f}% of all identified issues). This complete analysis covers all {num_trades} trade categories identified during the inspection."""
            
            summary_para = doc.add_paragraph(summary_text)
            summary_para.style = 'CleanBody'
    
    except Exception as e:
        print(f"Error creating pie chart: {e}")


def create_severity_chart(doc, metrics):
    """Create severity distribution chart"""
    
    if not MATPLOTLIB_AVAILABLE:
        add_text_severity_summary(doc, metrics)
        return
    
    try:
        chart_title = doc.add_paragraph("Unit Classification by Defect Severity")
        chart_title.style = 'CleanSubsectionHeader'
        
        if 'summary_unit' in metrics and len(metrics['summary_unit']) > 0:
            fig, ax = plt.subplots(figsize=(12, 7))
            
            units_data = metrics['summary_unit']
            
            categories = []
            counts = []
            colors = []
            
            # Calculate categories
            extensive_count = len(units_data[units_data['DefectCount'] >= 15])
            categories.append('Extensive\n(15+ defects)')
            counts.append(extensive_count)
            colors.append('#ff9999')
            
            major_count = len(units_data[(units_data['DefectCount'] >= 8) & (units_data['DefectCount'] <= 14)])
            categories.append('Major\n(8-14 defects)')
            counts.append(major_count)
            colors.append('#ffcc99')
            
            minor_count = len(units_data[(units_data['DefectCount'] >= 3) & (units_data['DefectCount'] <= 7)])
            categories.append('Minor\n(3-7 defects)')
            counts.append(minor_count)
            colors.append('#ffff99')
            
            ready_count = len(units_data[units_data['DefectCount'] <= 2])
            categories.append('Ready\n(0-2 defects)')
            counts.append(ready_count)
            colors.append('#99ff99')
            
            bars = ax.bar(categories, counts, color=colors, alpha=0.8)
            
            ax.set_ylabel('Number of Units', fontsize=14, fontweight='600')
            ax.set_title('Unit Distribution by Defect Severity Level', 
                        fontsize=16, fontweight='600', pad=20)
            ax.grid(axis='y', alpha=0.3, linestyle=':')
            
            # Value labels
            for bar, value in zip(bars, counts):
                if value > 0:
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(counts)*0.01,
                           f'{value}', ha='center', va='bottom', 
                           fontweight='bold', fontsize=12)
            
            plt.tight_layout()
            add_chart_to_document(doc, fig)
            plt.close()
    
    except Exception as e:
        print(f"Error creating severity chart: {e}")


def create_trade_chart(doc, metrics):
    """Create trade performance chart"""
    
    if not MATPLOTLIB_AVAILABLE:
        return
    
    try:
        trade_header = doc.add_paragraph("Trade Category Performance Analysis")
        trade_header.style = 'CleanSubsectionHeader'
        
        if 'summary_trade' not in metrics or len(metrics['summary_trade']) == 0:
            return
        
        top_trades = metrics['summary_trade'].head(10)
        
        fig, ax = plt.subplots(figsize=(12, 8))
        
        colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99', '#ff99cc'] * 2
        colors = colors[:len(top_trades)]
        
        if NUMPY_AVAILABLE:
            y_pos = np.arange(len(top_trades))
        else:
            y_pos = list(range(len(top_trades)))
        
        # Sanitize trade names
        trade_labels = [sanitize_text(str(trade)) for trade in top_trades['Trade']]
        
        bars = ax.barh(y_pos, top_trades['DefectCount'], color=colors, alpha=0.8)
        
        ax.set_yticks(y_pos)
        ax.set_yticklabels(trade_labels, fontsize=12)
        ax.set_xlabel('Number of Defects', fontsize=14, fontweight='600')
        ax.set_title('Trade Categories Ranked by Defect Frequency', 
                    fontsize=16, fontweight='600', pad=20)
        
        ax.grid(axis='x', alpha=0.3, linestyle=':')
        
        # Value labels
        total_defects = metrics.get('total_defects', 1)
        for i, (bar, value) in enumerate(zip(bars, top_trades['DefectCount'])):
            percentage = (value / total_defects * 100) if total_defects > 0 else 0
            ax.text(bar.get_width() + max(top_trades['DefectCount']) * 0.02, 
                   bar.get_y() + bar.get_height()/2,
                   f'{value} ({percentage:.1f}%)', va='center', 
                   fontweight='600', fontsize=10)
        
        plt.tight_layout()
        add_chart_to_document(doc, fig)
        plt.close()
    
    except Exception as e:
        print(f"Error creating trade chart: {e}")


def add_chart_to_document(doc, fig):
    """Helper to add matplotlib charts to document"""
    
    try:
        chart_buffer = BytesIO()
        fig.savefig(chart_buffer, format='png', dpi=300, bbox_inches='tight', 
                    facecolor='white', edgecolor='none', pad_inches=0.2)
        chart_buffer.seek(0)
        
        chart_para = doc.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart_run = chart_para.add_run()
        chart_run.add_picture(chart_buffer, width=Inches(7))
        
        doc.add_paragraph()
    
    except Exception as e:
        print(f"Error adding chart: {e}")


# ═══════════════════════════════════════════════════════════════════
# TEXT SUMMARIES (when matplotlib not available)
# ═══════════════════════════════════════════════════════════════════

def add_text_trade_summary(doc, metrics):
    """Text-based trade summary"""
    try:
        breakdown_header = doc.add_paragraph("Defects Distribution by Trade Category")
        breakdown_header.style = 'CleanSubsectionHeader'
        
        note_para = doc.add_paragraph("(Visual charts require matplotlib - showing text summary)")
        note_para.style = 'CleanBody'
        
        if 'summary_trade' not in metrics or len(metrics['summary_trade']) == 0:
            return
        
        trade_data = metrics['summary_trade'].copy()
        total_defects = metrics.get('total_defects', 0)
        
        if total_defects > 0:
            for idx, (_, row) in enumerate(trade_data.iterrows(), 1):
                percentage = (row['DefectCount'] / total_defects * 100)
                trade_text = f"{idx}. {sanitize_text(row['Trade'])}: {row['DefectCount']} defects ({percentage:.1f}%)"
                trade_para = doc.add_paragraph(trade_text)
                trade_para.style = 'CleanBody'
                trade_para.paragraph_format.left_indent = Inches(0.3)
    
    except Exception as e:
        print(f"Error in text trade summary: {e}")


def add_text_severity_summary(doc, metrics):
    """Text-based severity summary"""
    try:
        chart_title = doc.add_paragraph("Unit Classification by Defect Severity")
        chart_title.style = 'CleanSubsectionHeader'
        
        note_para = doc.add_paragraph("(Visual charts require matplotlib - showing text summary)")
        note_para.style = 'CleanBody'
        
        if 'summary_unit' in metrics and len(metrics['summary_unit']) > 0:
            units_data = metrics['summary_unit']
            
            extensive_count = len(units_data[units_data['DefectCount'] >= 15])
            major_count = len(units_data[(units_data['DefectCount'] >= 8) & (units_data['DefectCount'] <= 14)])
            minor_count = len(units_data[(units_data['DefectCount'] >= 3) & (units_data['DefectCount'] <= 7)])
            ready_count = len(units_data[units_data['DefectCount'] <= 2])
            
            severity_data = [
                ("Extensive (15+ defects)", extensive_count),
                ("Major (8-14 defects)", major_count),
                ("Minor (3-7 defects)", minor_count),
                ("Ready (0-2 defects)", ready_count)
            ]
            
            for category, count in severity_data:
                if count > 0:
                    severity_text = f"• {category}: {count} units"
                    severity_para = doc.add_paragraph(severity_text)
                    severity_para.style = 'CleanBody'
                    severity_para.paragraph_format.left_indent = Inches(0.3)
    
    except Exception as e:
        print(f"Error in text severity summary: {e}")


def add_text_units_summary(doc, metrics):
    """Text-based units summary"""
    try:
        chart_title = doc.add_paragraph("Top 20 Units Requiring Immediate Intervention")
        chart_title.style = 'CleanSubsectionHeader'
        
        note_para = doc.add_paragraph("(Visual charts require matplotlib - showing text summary)")
        note_para.style = 'CleanBody'
        
        if 'summary_unit' not in metrics or len(metrics['summary_unit']) == 0:
            return
        
        top_units = metrics['summary_unit'].head(20)
        
        for idx, (_, row) in enumerate(top_units.iterrows(), 1):
            unit_text = f"{idx}. Unit {sanitize_text(str(row['Unit']))}: {row['DefectCount']} defects"
            unit_para = doc.add_paragraph(unit_text)
            unit_para.style = 'CleanBody'
            unit_para.paragraph_format.left_indent = Inches(0.3)
    
    except Exception as e:
        print(f"Error in text units summary: {e}")


# ═══════════════════════════════════════════════════════════════════
# DATA PROCESSING HELPERS
# ═══════════════════════════════════════════════════════════════════

def generate_complete_component_details(processed_data):
    """Generate component details for trade tables - FIXED"""
    
    try:
        print("🔍 Generating component details...")
        print(f"   Input shape: {processed_data.shape}")
        print(f"   Columns: {list(processed_data.columns)}")
        
        # Check required columns
        required_columns = ['StatusClass', 'Trade', 'Room', 'Component', 'Unit']
        missing_columns = [col for col in required_columns if col not in processed_data.columns]
        
        if missing_columns:
            print(f"   ❌ Missing columns: {missing_columns}")
            return pd.DataFrame()
        
        # Filter for defects - StatusClass should already be "Not OK"
        defects_only = processed_data[processed_data['StatusClass'] == 'Not OK'].copy()
        print(f"   Defects found: {len(defects_only)}")
        
        if len(defects_only) == 0:
            print("   ❌ No defects to process")
            return pd.DataFrame()
        
        # Group by trade, room, component
        print("   Grouping data...")
        component_summary = defects_only.groupby(['Trade', 'Room', 'Component']).agg({
            'Unit': lambda x: ', '.join(sorted([str(u) for u in x.unique()]))
        }).reset_index()
        
        component_summary.columns = ['Trade', 'Room', 'Component', 'Affected Units']
        
        # Count unique units
        unit_counts = defects_only.groupby(['Trade', 'Room', 'Component'])['Unit'].nunique().reset_index()
        component_summary = component_summary.merge(unit_counts, on=['Trade', 'Room', 'Component'])
        component_summary.columns = ['Trade', 'Room', 'Component', 'Affected Units', 'Unit Count']
        
        # Sort
        component_summary = component_summary.sort_values(['Trade', 'Unit Count'], ascending=[True, False])
        
        print(f"   ✅ Generated {len(component_summary)} component rows")
        print(f"   Sample:\n{component_summary.head()}")
        
        return component_summary
    
    except Exception as e:
        print(f"   ❌ Error generating component details: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()


def generate_component_breakdown(processed_data):
    """Generate component breakdown - FIXED"""
    
    try:
        print("🔍 Generating component breakdown...")
        
        required_columns = ['StatusClass', 'Trade', 'Component', 'Unit']
        missing_columns = [col for col in required_columns if col not in processed_data.columns]
        
        if missing_columns:
            print(f"   ❌ Missing columns: {missing_columns}")
            return pd.DataFrame()
        
        # Filter for defects
        defects_only = processed_data[processed_data['StatusClass'] == 'Not OK'].copy()
        print(f"   Defects found: {len(defects_only)}")
        
        if len(defects_only) == 0:
            return pd.DataFrame()
        
        # Group by component and trade
        component_summary = defects_only.groupby(['Component', 'Trade']).agg({
            'Unit': lambda x: ', '.join(sorted([str(u) for u in x.unique()]))
        }).reset_index()
        
        component_summary.columns = ['Component', 'Trade', 'Affected_Units']
        
        # Count unique units
        unit_counts = defects_only.groupby(['Component', 'Trade'])['Unit'].nunique().reset_index()
        component_summary = component_summary.merge(unit_counts, on=['Component', 'Trade'])
        component_summary.columns = ['Component', 'Trade', 'Affected_Units', 'Unit_Count']
        
        # Sort
        component_summary = component_summary.sort_values(['Unit_Count', 'Component'], ascending=[False, True])
        
        print(f"   ✅ Generated {len(component_summary)} component rows")
        
        return component_summary
    
    except Exception as e:
        print(f"   ❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return pd.DataFrame()

def add_trade_tables(doc, component_details):
    """Add trade-specific tables with shading"""
    
    try:
        if len(component_details) == 0:
            return
        
        trades = component_details['Trade'].unique()
        
        for trade in trades:
            try:
                trade_data = component_details[component_details['Trade'] == trade]
                
                trade_header = doc.add_paragraph(sanitize_text(str(trade)))
                trade_header.style = 'CleanSubsectionHeader'
                
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                table.columns[0].width = Inches(2.5)
                table.columns[1].width = Inches(4.0)
                table.columns[2].width = Inches(0.8)
                
                # Headers
                headers = ['Component & Location', 'Affected Units', 'Count']
                for i, header in enumerate(headers):
                    cell = table.cell(0, i)
                    cell.text = header
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.runs[0]
                    run.font.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    set_cell_background_color(cell, "F0F0F0")
                
                # Data rows with alternating colors
                for idx, (_, row) in enumerate(trade_data.iterrows()):
                    table_row = table.add_row()
                    row_color = "FFFFFF" if idx % 2 == 0 else "F8F8F8"
                    
                    # Component & Location
                    component_location = sanitize_text(str(row['Component']))
                    if pd.notna(row['Room']) and str(row['Room']).strip():
                        component_location += f" ({sanitize_text(str(row['Room']))})"
                    
                    cell1 = table_row.cells[0]
                    cell1.text = component_location
                    cell1.paragraphs[0].runs[0].font.name = 'Arial'
                    cell1.paragraphs[0].runs[0].font.size = Pt(10)
                    cell1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    set_cell_background_color(cell1, row_color)
                    
                    # Affected Units
                    cell2 = table_row.cells[1]
                    cell2.text = sanitize_text(str(row['Affected Units']))
                    cell2.paragraphs[0].runs[0].font.name = 'Arial'
                    cell2.paragraphs[0].runs[0].font.size = Pt(10)
                    cell2.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    set_cell_background_color(cell2, row_color)
                    
                    # Count
                    cell3 = table_row.cells[2]
                    cell3.text = str(row['Unit Count'])
                    cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell3.paragraphs[0].runs[0].font.name = 'Arial'
                    cell3.paragraphs[0].runs[0].font.size = Pt(10)
                    cell3.paragraphs[0].runs[0].font.bold = True
                    cell3.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    set_cell_background_color(cell3, row_color)
                
                doc.add_paragraph()
            
            except Exception as e:
                print(f"Error processing trade {trade}: {e}")
                continue
    
    except Exception as e:
        print(f"Error in trade tables: {e}")


def set_cell_background_color(cell, color_hex):
    """Set cell background color"""
    
    try:
        shading_elm = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
    except Exception as e:
        print(f"Could not set cell background: {e}")

def create_error_document(error, metrics):
    """Error document"""
    doc = Document()
    doc.add_heading("Inspection Report - Generation Error", level=1)
    doc.add_paragraph(f"Error: {str(error)}")
    return doc

# ═══════════════════════════════════════════════════════════════════
# MAIN EXECUTION
# ═══════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("✅ Professional Word Generator - Database/API Version")
    print("📋 Based on working CSV template")
    print("🎨 Features: Arial font, professional formatting, &amp; fix")
    print(f"📊 Matplotlib: {'Available' if MATPLOTLIB_AVAILABLE else 'Not Available'}")
    print(f"🔢 NumPy: {'Available' if NUMPY_AVAILABLE else 'Not Available'}")


